import os
import shutil
from pypdf import PdfReader
from docx2pdf import convert
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, RGBColor, Cm, Pt
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def extract_raw_document_images(fileName):
    fileName = fileName.replace(".docx", "")
    convert(f"assets/{fileName}.docx", f"assets/{fileName}.pdf")
    pdfPath = f"assets/{fileName}.pdf"
    path = "assets/images"
    
    if os.path.exists(path) and os.path.isdir(path):
        shutil.rmtree(path)

    os.makedirs("assets/images")

    reader = PdfReader(f"assets/{fileName}.pdf")
    page = reader.pages[2]
    count = 0
    images = []
    imageNames = []
    count = 0
    i = 0
    for page in reader.pages:
        for img in page.images:
            if i != 1 and img.name not in imageNames:
                count += 1
                images.append(img)
                imageNames.append(img.name)
                with open("assets/images/" + str(count) + '.' + img.name.split('.')[1],"wb") as fp:
                    fp.write(img.data)
        i+=1
            
def copy_paragraph(output_doc, paragraph):

    output_paragraph = output_doc.add_paragraph()
    # Alignment data of whole paragraph
    output_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
    i=0
    for row in paragraph.runs:
        output_row = output_paragraph.add_run(row.text)
        # Font data
        output_row.style.name = "Normal"
        # Size of font data
        if row.font.size != None:
            output_row.font.size = row.font.size-1000
        else:
            output_row.font.size = row.font.size-1000
        # Bold data
        output_row.bold = row.bold
        # Italic data
        output_row.italic = row.italic
        # Underline data
        output_row.underline = row.underline
        # Color data
        output_row.font.color.rgb = row.font.color.rgb

def add_h1(output_doc, text):
    heading = output_doc.add_heading(text, 1)
    heading.style.font.bold = True
    heading.style.font.name = 'Cambria'
    heading.style.font.size = 250000
    heading.style.font.color.rgb = RGBColor(250, 168, 32)
       
def add_h2(output_doc, text):
    heading = output_doc.add_heading(text, 2)
    heading.style.font.name = 'Cambria'
    heading.style.font.size = 200000
    heading.style.font.bold = False
    heading.line_spacing_rule = WD_LINE_SPACING.SINGLE
    heading.style.font.color.rgb = RGBColor(250, 168, 32)
    
def add_h3(output_doc, text):
    heading = output_doc.add_heading(text, 3)
    heading.style.font.name = 'Cambria'
    heading.style.font.size = 150000
    heading.style.font.bold = False
    heading.style.font.color.rgb = RGBColor(128, 128, 128)

def format_table(output_doc):
    width = (Inches(4.5), Inches(4.5), Inches(1.5))
    t = 0
    for table in output_doc.tables:
        table.style = "Table Grid"
        i = 0
        for r in table.rows:
            j=0
            if len(table.rows[i].cells) == 2:
                width = (Inches(6), Inches(6))
            elif len(table.rows[i].cells) == 3:
                width = (Inches(6.5), Inches(4.5), Inches(2))
            elif len(table.rows[i].cells) == 4:
                width = (Inches(4), Inches(3), Inches(1), Inches(4))
            elif len(table.rows[i].cells) == 5:
                width = (Inches(2), Inches(4), Inches(3), Inches(2), Inches(2))
            elif len(table.rows[i].cells) == 6:
                width = (Inches(2), Inches(2), Inches(2), Inches(2), Inches(2), Inches(2))
            elif len(table.rows[i].cells) == 7:
                width = (Inches(0.5), Inches(2), Inches(2), Inches(2), Inches(2.5), Inches(1.5), Inches(1.5))
            for cell in table.rows[i].cells:
                table.rows[i].cells[j].width = width[j]
                j+=1
            i+=1
        t+=1

def format_table_with_picture(output_doc, tableNo, imagePath):
    output_doc.tables[tableNo].style = "Table Grid"
    paragraph = output_doc.tables[tableNo].rows[0].cells[3].paragraphs[0]
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
    paragraph = output_doc.tables[tableNo].rows[0].cells[3].add_paragraph()
    run = paragraph.add_run()
    run.add_picture(imagePath, width=Cm(4.5), height=Cm(4.5))

def copy_table(output_doc, table):
    p = output_doc.paragraphs[-1]
    new_tbl = deepcopy(table._tbl)
    p._p.addnext(new_tbl)

def title_run(r):
    r.font.size = Pt(25.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(250, 168, 32)
    r.font.name = "Cambria (Heading)"
    
def replace_variables(output_doc, raw):
    txbx = raw.inline_shapes._body.xpath('//w:txbxContent')
    for tx_idx, tx in enumerate(txbx):
            children = tx.getchildren()
            i=0
            for child_idx, child in enumerate(children):
                if child.text != "":
                    if child.text.startswith("Angebotsnr."):
                        id = child.text
                    elif i==0:
                        address = child.text
                        i+=1
    
    module = raw.tables[1].cell(4,1).paragraphs[0].text
    kw = raw.tables[1].cell(2,1).paragraphs[0].text
    date = raw.paragraphs[0].text
    address += ', ' + raw.tables[0].cell(1,0).paragraphs[0].text
    try:
        address += ', ' + raw.tables[0].cell(1,0).paragraphs[1].text
    except:
        print('')
    address = address.split(', ')
    
    i=0
    for p in output_doc.paragraphs:
        text = p.text
        if i==2:
            p.text = ""
            r = p.add_run(text.replace("0.00",kw))
            title_run(r)
        if i==3:
            p.text = ""
            r = p.add_run(text.replace("0",module))
            title_run(r)
        elif i == 5:
            p.text = ""
            r = p.add_run(address[0])
            title_run(r)
        elif i == 6:
            p.text = ""
            if len(address) > 2:
                r = p.add_run(address[2])
                title_run(r)
        elif i == 7:
            p.text = ""
            r = p.add_run(address[1])
            title_run(r)
        elif i==8:
            p.text = ""
            r = p.add_run(date)
            title_run(r)
            
        i+=1

def prepare_header(output_doc, raw_doc):
    txbx = raw_doc.inline_shapes._body.xpath('//w:txbxContent')
    for tx_idx, tx in enumerate(txbx):
            children = tx.getchildren()
            i=0
            for child_idx, child in enumerate(children):
                if child.text != "":
                    if child.text.startswith("Angebotsnr."):
                        id = child.text
                    elif i==0:
                        address = child.text
                        i+=1
                    
    date = raw_doc.paragraphs[0].text
    address += ', ' + raw_doc.tables[0].cell(1,0).paragraphs[0].text
    try:
        address += ', ' + raw.tables[0].cell(1,0).paragraphs[1].text
    except:
        print('')
    address = address.split(', ')
    
    section = output_doc.sections[0]
    header = section.header
    section.different_first_page_header_footer = True
    t = header.add_table(2,2, Inches(24))

    t.cell(0,0).text = ''
    t.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = t.cell(0,0).paragraphs[0].add_run(address[0])
    r1.font.size = Pt(9)
    r1.add_break()
    if len(address) > 2:
        r2 = t.cell(0,0).paragraphs[0].add_run(address[2])
        r2.font.size = Pt(9)
        r2.add_break()
    r3 = t.cell(0,0).paragraphs[0].add_run(address[1])
    r3.font.size = Pt(9)

    t.cell(1,0).text = ''
    t.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r11 = t.cell(1,0).paragraphs[0].add_run(date)
    r11.font.size = Pt(9)
    r11.add_break()
    r21 = t.cell(1,0).paragraphs[0].add_run(id.split(" ")[1])
    r21.font.size = Pt(9)

    t.cell(0,1).text = ''
    t.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r21 = t.cell(0,1).paragraphs[0].add_run()
    r21.add_picture("assets/template_images/header.png", Cm(1.5), Cm(1.25))

    t.cell(1,1).text = ''
    t.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r31 = t.cell(1,1).paragraphs[0].add_run("Solardach 24 GmbH")
    r31.font.color.rgb = RGBColor(250, 168, 32)
    r31.font.size = Pt(9)
    r31.add_break()
    r42 = t.cell(1,1).paragraphs[0].add_run("Sicher und zuverlässig")
    r42.font.color.rgb = RGBColor(250, 168, 32)
    r42.font.size = Pt(9)

def prepare_footer(output_doc):
    footer = output_doc.sections[0].footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("Solardach24 GmbH")
    r1.add_break()
    r2 = p.add_run("Reinacherstrasse 261 ∙ 4053 Basel")
    r2.add_break()
    r3=p.add_run("Collègegasse 9 ∙ 2502 Biel/Bienne")
    r3.add_break()
    r3=p.add_run("+41 61 511 22 22 ∙ office@solardach24.ch ∙ CHE-152.292-000")
    
def add_toc(output_doc):
    para = output_doc.add_paragraph("INHALTSVERZEICHNIS")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.bold = True
        run.underline = True
    output_doc.add_paragraph(" ")
    paragraph = output_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar3 = OxmlElement('w:updateFields') 
    fldChar3.set(qn('w:val'), 'true') 
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(250, 168, 32)
    
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    p_element = paragraph._p
    
def add_picture_inline(output_doc, picture_path, width, height):
    try:
        output_doc.add_picture(picture_path + ".png", width, height)
    except:
        try: 
            output_doc.add_picture(picture_path + ".jpg", width, height)
        except:
            print(f"log: {picture_path} not found")

def extract_para_style(raw, para_style):
    para = []
    for p in raw.paragraphs:
        if p.style.name.startswith(para_style):
                para.append(p.text)
    return para
                
                
if __name__=="__main__":
    
    print("Enter Raw File Name: ")
    fileName = input()
    fileName = fileName.replace(".docx", "")
    template = Document('assets/template.docx')
    raw = Document(f'assets/{fileName}.docx')
    doc = Document()
    print(f"Reading file assets/{fileName}.docx")
    print(f"Total number of tables: {len(raw.tables)}")
    
    # Extract images from raw document - these a document specific images
    extract_raw_document_images(fileName)

    doc.add_picture("assets/template_images/image1.png", width=Inches(6), height=Inches(4))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ## headings upto table of contents
    doc.add_paragraph(" ")
    for paragraph in template.paragraphs[12:20]:
        copy_paragraph(doc, paragraph)

    # -------------------------------------------------  
    doc.add_paragraph(" ")
    add_toc(doc)
    doc.add_page_break()
    # ---------------------------------------------------

    tableIndex = 1
    h1_index = 0
    h2_index = 0
    h3_index = 0
    pic_index = 2
    h1 = extract_para_style(raw, 'Heading 1')
    h2 = extract_para_style(raw, 'Heading 2')
    h3 = extract_para_style(raw, 'Heading 3')
    para = []
    for p in raw.paragraphs:
        if p.text != '':
            para.append(p.text)
    
    add_h1(doc, f"1. {h1[h1_index]}")
    h1_index+=1
    add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
    pic_index+=1
    
    doc.add_paragraph(" ")
    if "PV-Anlage" in h2 :
        add_h2(doc, "PV-Anlage")
        
        i = para.index('PV-Anlage')
        add_h3(doc, para[i+1])
        para.remove('PV-Anlage')
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1
        add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
        if pic_index != 3:
            pic_index += 1
        pic_index+=1
    
    if "Ertragsprognose" in h2 :
        add_h2(doc, "Ertragsprognose")
        add_h3(doc, "Ertragsprognose")
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1

    # ----------------------------------------
    doc.add_page_break()
    # ----------------------------------------

    add_h1(doc, f"2. {h1[h1_index]}")
    h1_index+=1
    
    if "Überblick" in h2 :
        add_h2(doc, "Überblick")

        add_h3(doc, "Anlagendaten")
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1
        
        add_h3(doc, "Klimadaten")
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1
        
        add_h3(doc, "Verbrauch")
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1

        doc.add_paragraph("")
        add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
        pic_index+=1

    if "Modulflächen" in h2 :
        add_h2(doc, "Modulflächen")
        
        i = para.index("Modulflächen")

        while True:
            if 'Modulfläche' in para[i+1]:
                add_h2(doc, para[i+1])
                add_h3(doc, para[i+2])
                
                copy_table(doc, raw.tables[tableIndex])
                # # print(tableIndex)
                tableIndex+=1
                doc.add_paragraph("")
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
                pic_index+=1
                i+=4
            else:
                break

    add_h2(doc, "Horizontlinie, 3D-Planung")
    add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
    pic_index+=1

    # ----------------------------------------
    doc.add_page_break()
    # ----------------------------------------

    if "Wechselrichterverschaltung" in h2 :
        add_h2(doc, "Wechselrichterverschaltung")
        i = para.index("Wechselrichterverschaltung")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # # print(tableIndex)
        tableIndex+=1

    if "AC-Netz" in h2 :
        add_h2(doc, "AC-Netz")
        i = para.index("AC-Netz")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # # print(tableIndex)
        tableIndex+=1

    if "Batteriesysteme" in h2 :
        add_h2(doc, "Batteriesysteme")
        i = para.index("Batteriesysteme")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1

    # ----------------------------------------
    doc.add_page_break()
    # ----------------------------------------

    add_h1(doc, f"3. {h1[h1_index]}")
    h1_index+=1
    
    if "Ergebnisse Gesamtanlage" in h2 :
        add_h2(doc, "Ergebnisse Gesamtanlage")
        i = para.index("Ergebnisse Gesamtanlage")
        i+=1
        if "PV-Anlage" in para:
            add_h3(doc, "PV-Anlage")
            # here index represent table no of the raw document
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            # index represent table no of the output document
            try:
                format_table_with_picture(doc, tableIndex-1, f"assets/images/{pic_index}.png")
            except:
                format_table_with_picture(doc, tableIndex-1, f"assets/images/{pic_index}.jpg")
            pic_index+=1
            tableIndex+=1
            i+=1
        
        if "Verbraucher" in para:
            add_h3(doc, "Verbraucher")
            # here index represent table no of the raw document
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)

            # index represent table no of the output document
            try:
                format_table_with_picture(doc, tableIndex-1, f"assets/images/{pic_index}.png")
            except: 
                format_table_with_picture(doc, tableIndex-1, f"assets/images/{pic_index}.jpg")
            pic_index+=1
            tableIndex+=1
            i+=1
        
        if "Batteriesystem" in para:
            add_h3(doc, "Batteriesystem")
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1
            i+=1

        if "Autarkiegrad" in para:
            add_h3(doc, "Autarkiegrad")
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1
            i+=1

        doc.add_paragraph("")

        while True:
            if 'Abbildung' in para[i]: 
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
                pic_index+=1
                i+=1
            else: 
                break

    if "Ergebnisse pro Modulfläche" in h2 :
        add_h2(doc, "Ergebnisse pro Modulfläche")
        i = para.index("Ergebnisse pro Modulfläche")
        
        if para[i+1] != h1[h1_index]:
            add_h3(doc, para[i+1])
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1
            i+=1

        if para[i+1] != h1[h1_index]:
            add_h3(doc, para[i+1])
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1

    add_h1(doc, f"4. {h1[h1_index]}")
    h1_index+=1
    
    if "Datenblatt PV-Modul" in h2 :
        add_h2(doc, "Datenblatt PV-Modul")
        i = para.index("Datenblatt PV-Modul")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1

    if "Datenblatt Wechselrichter" in h2 :
        add_h2(doc, "Datenblatt Wechselrichter")
        i = para.index("Datenblatt Wechselrichter")
        
        if para[i+1].startswith("Wechselrichter:"):
            add_h3(doc, para[i+1])
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1
            i+=1

        if para[i+1].startswith("Wechselrichter:"):
            add_h3(doc, para[i+1])
            copy_table(doc, raw.tables[tableIndex])
            # print(tableIndex)
            tableIndex+=1

    if "Datenblatt Batteriesystem" in h2 :
        add_h2(doc, "Datenblatt Batteriesystem")
        i = para.index("Datenblatt Batteriesystem")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1

    if "Datenblatt Batterie" in h2 :
        add_h2(doc, "Datenblatt Batterie")
        i = para.index("Datenblatt Batterie")
        add_h3(doc, para[i+1])
        # print(tableIndex)
        copy_table(doc, raw.tables[tableIndex])
        tableIndex+=1

    # ----------------------------------------
    doc.add_page_break()
    # ----------------------------------------

    add_h1(doc, f"5. {h1[h1_index]}")
    h1_index+=1

    if "Schaltplan" in h2 :
        add_h2(doc, "Schaltplan")
        i = para.index("Schaltplan")
        while True:
            if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                pic_index+=1
                i+=1
            else: 
                break

    if "Übersichtsplan" in h2 :
        add_h2(doc, "Übersichtsplan")
        i = para.index("Übersichtsplan")
        while True:
            if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                pic_index+=1
                i+=1
            else: 
                break

    if "Bemaßungsplan" in h2 :
        add_h2(doc, "Bemaßungsplan")
        i = para.index("Bemaßungsplan")
        while True:
            if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                pic_index+=1
                i+=1
            else: 
                break

    if "Strangplan" in h2 :
        add_h2(doc, "Strangplan")
        i = para.index("Strangplan")
        while True:
            if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                pic_index+=1
                i+=1
            else: 
                break

    if "Stückliste" in h2 :
        add_h2(doc, "Stückliste")
        i = para.index("Stückliste")
        add_h3(doc, para[i+1])
        copy_table(doc, raw.tables[tableIndex])
        # print(tableIndex)
        tableIndex+=1
        
    # ----------------------------------------
    doc.add_page_break()
    # ----------------------------------------   
    
    add_h1(doc, "6. Warum Solardach24 GmbH?")
    doc.add_picture("assets/template_images/image2.png", width=Inches(6.5), height=Inches(7))

    add_h1(doc, "7. Vier Köpfe. Für Ihre PV-Anlage.")
    doc.add_picture("assets/template_images/image3.png", width=Inches(6.5), height=Inches(7))

    add_h1(doc, "8. Unser Partner. Für Ihre Sicherheit.")
    doc.add_picture("assets/template_images/image4.png", width=Inches(6.5), height=Inches(7))

    add_h1(doc, "9. Unser Haustechnik-Partner. Für Ihre persönliche Energiewende.")
    doc.add_picture("assets/template_images/image5.png", width=Inches(6.5), height=Inches(6))
    doc.add_picture("assets/template_images/image6.png", width=Inches(6.5), height=Inches(7.5))
    doc.add_picture("assets/template_images/image7.png", width=Inches(6.5), height=Inches(7.5))

    prepare_header(doc, raw)
    prepare_footer(doc)
    replace_variables(doc, raw)
    format_table(doc)

    doc.save(f'assets/{fileName}-output.docx')
    print(f'assets/{fileName}-output.docx created')
